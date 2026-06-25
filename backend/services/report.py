"""
Report Generation Service
- สร้าง PDF รายงานการประชุม (ReportLab)
- สร้าง DOCX รายงานการประชุม (python-docx)
- หัวจดหมาย ยสท. format ทางการ
"""
import os
from datetime import datetime
from typing import Optional

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---- ลงทะเบียน font ภาษาไทย สำหรับ ReportLab ----
_FONT_REGISTERED = False

def _register_thai_font():
    global _FONT_REGISTERED
    if _FONT_REGISTERED:
        return
    # ใช้ Sarabun จาก Google Fonts (ต้องมีไฟล์ .ttf ใน fonts/)
    font_dir = os.path.join(os.path.dirname(__file__), "..", "fonts")
    sarabun_regular = os.path.join(font_dir, "Sarabun-Regular.ttf")
    sarabun_bold = os.path.join(font_dir, "Sarabun-Bold.ttf")

    if os.path.exists(sarabun_regular):
        pdfmetrics.registerFont(TTFont("Sarabun", sarabun_regular))
    if os.path.exists(sarabun_bold):
        pdfmetrics.registerFont(TTFont("Sarabun-Bold", sarabun_bold))

    _FONT_REGISTERED = True


def _thai(text: str, size: int = 12, bold: bool = False) -> Paragraph:
    """สร้าง Paragraph ภาษาไทย"""
    _register_thai_font()
    font = "Sarabun-Bold" if bold else "Sarabun"
    style = ParagraphStyle(
        name="thai",
        fontName=font,
        fontSize=size,
        leading=size * 1.6,
        wordWrap="CJK",
    )
    return Paragraph(text, style)


# ===== PDF =====

def generate_pdf(
    output_path: str,
    meeting_title: str,
    meeting_date: str,
    meeting_location: str,
    attendees: list[str],
    agenda_items: list[str],
    summary_by_agenda: list[dict],
    action_items: list[dict],
    transcript_text: Optional[str] = None,
) -> str:
    """สร้างไฟล์ PDF รายงานการประชุม คืนค่า path ของไฟล์"""
    _register_thai_font()

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        leftMargin=3 * cm,
        rightMargin=2.5 * cm,
    )

    story = []
    GREEN = colors.HexColor("#1B4332")

    # ---- หัวจดหมาย ----
    story.append(_thai("การยาสูบแห่งประเทศไทย (ยสท.)", 16, bold=True))
    story.append(_thai("สำนักบริหารทรัพย์สิน", 13))
    story.append(Spacer(1, 0.3 * cm))
    story.append(HRFlowable(width="100%", thickness=2, color=GREEN))
    story.append(Spacer(1, 0.3 * cm))

    # ---- ชื่อการประชุม ----
    story.append(_thai(f"รายงานการประชุม", 14, bold=True))
    story.append(_thai(meeting_title, 13, bold=True))
    story.append(Spacer(1, 0.4 * cm))

    # ---- ข้อมูลการประชุม ----
    info_data = [
        ["วันที่:", meeting_date],
        ["สถานที่:", meeting_location],
        ["จำนวนผู้เข้าร่วม:", f"{len(attendees)} ท่าน"],
    ]
    _register_thai_font()
    info_table = Table(info_data, colWidths=[4 * cm, 12 * cm])
    info_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Sarabun"),
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("FONTNAME", (0, 0), (0, -1), "Sarabun-Bold"),
        ("TEXTCOLOR", (0, 0), (0, -1), GREEN),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 0.5 * cm))

    # ---- ผู้เข้าร่วม ----
    if attendees:
        story.append(_thai("ผู้เข้าร่วมประชุม", 12, bold=True))
        story.append(Spacer(1, 0.2 * cm))
        for name in attendees:
            story.append(_thai(f"• {name}", 11))
        story.append(Spacer(1, 0.5 * cm))

    # ---- สรุปตามวาระ ----
    story.append(HRFlowable(width="100%", thickness=1, color=GREEN))
    story.append(Spacer(1, 0.3 * cm))
    story.append(_thai("สรุปการประชุมตามวาระ", 13, bold=True))
    story.append(Spacer(1, 0.3 * cm))

    for agenda_summary in summary_by_agenda:
        n = agenda_summary["agenda_number"]
        title = agenda_summary["agenda_title"]
        summary = agenda_summary["summary"]

        story.append(_thai(f"วาระที่ {n}: {title}", 12, bold=True))
        story.append(Spacer(1, 0.15 * cm))

        # แยกบรรทัดจาก summary
        for line in summary.split("\n"):
            line = line.strip()
            if not line or line.startswith("##"):
                continue
            if line.startswith("**") and line.endswith("**"):
                story.append(_thai(line.replace("**", ""), 11, bold=True))
            elif line.startswith(("-", "•")):
                story.append(_thai(f"  {line}", 11))
            else:
                story.append(_thai(line, 11))
        story.append(Spacer(1, 0.4 * cm))

    # ---- Action Items ----
    if action_items:
        story.append(HRFlowable(width="100%", thickness=1, color=GREEN))
        story.append(Spacer(1, 0.3 * cm))
        story.append(_thai("Action Items", 13, bold=True))
        story.append(Spacer(1, 0.2 * cm))
        for i, item in enumerate(action_items, 1):
            story.append(_thai(f"{i}. {item['text']}", 11))
        story.append(Spacer(1, 0.5 * cm))

    # ---- footer ----
    story.append(Spacer(1, 1 * cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
    story.append(Spacer(1, 0.2 * cm))
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
    story.append(_thai(f"จัดทำโดยระบบอัตโนมัติ ยสท. | {generated_at}", 9))

    doc.build(story)
    return output_path


# ===== DOCX =====

def generate_docx(
    output_path: str,
    meeting_title: str,
    meeting_date: str,
    meeting_location: str,
    attendees: list[str],
    agenda_items: list[str],
    summary_by_agenda: list[dict],
    action_items: list[dict],
) -> str:
    """สร้างไฟล์ DOCX รายงานการประชุม คืนค่า path ของไฟล์"""
    doc = Document()

    # ตั้งค่าหน้ากระดาษ A4
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    GREEN = RGBColor(0x1B, 0x43, 0x32)

    def add_heading(text: str, level: int = 1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = GREEN
        run.font.size = Pt(16 - (level - 1) * 2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def add_body(text: str, size: int = 11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Sarabun"
        return p

    # ---- หัวเรื่อง ----
    t = doc.add_paragraph()
    r = t.add_run("การยาสูบแห่งประเทศไทย (ยสท.)")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = GREEN
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    t2 = doc.add_paragraph()
    r2 = t2.add_run("สำนักบริหารทรัพย์สิน")
    r2.font.size = Pt(13)
    r2.font.color.rgb = GREEN
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_heading(f"รายงานการประชุม: {meeting_title}", level=1)

    # ---- ตารางข้อมูลการประชุม ----
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    cells = [
        ("วันที่", meeting_date),
        ("สถานที่", meeting_location),
        ("จำนวนผู้เข้าร่วม", f"{len(attendees)} ท่าน"),
    ]
    for i, (label, value) in enumerate(cells):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Sarabun"

    doc.add_paragraph()

    # ---- ผู้เข้าร่วม ----
    if attendees:
        add_heading("ผู้เข้าร่วมประชุม", level=2)
        for name in attendees:
            add_body(f"• {name}")
        doc.add_paragraph()

    # ---- สรุปตามวาระ ----
    add_heading("สรุปการประชุมตามวาระ", level=2)

    for ag in summary_by_agenda:
        p = doc.add_paragraph()
        r = p.add_run(f"วาระที่ {ag['agenda_number']}: {ag['agenda_title']}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = GREEN

        for line in ag["summary"].split("\n"):
            line = line.strip()
            if not line or line.startswith("##"):
                continue
            if line.startswith("**") and line.endswith("**"):
                p2 = doc.add_paragraph()
                r2 = p2.add_run(line.replace("**", ""))
                r2.bold = True
                r2.font.size = Pt(11)
            else:
                add_body(line)

        doc.add_paragraph()

    # ---- Action Items ----
    if action_items:
        add_heading("Action Items", level=2)
        for i, item in enumerate(action_items, 1):
            add_body(f"{i}. {item['text']}")

    doc.add_paragraph()
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
    footer_p = doc.add_paragraph()
    footer_r = footer_p.add_run(f"จัดทำโดยระบบอัตโนมัติ ยสท. | {generated_at}")
    footer_r.font.size = Pt(9)
    footer_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(output_path)
    return output_path
